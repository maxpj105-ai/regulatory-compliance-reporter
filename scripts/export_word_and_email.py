import os
import sys
import argparse
import datetime
import smtplib
import subprocess
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    """設定表格儲存格背景底色"""
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """設定儲存格邊距 (dxa)"""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_styled_docx(md_content, output_docx_path):
    doc = docx.Document()
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    lines = md_content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if line.startswith('# '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line[2:].strip())
            run.font.name = 'Microsoft JhengHei'
            run.font.size = Pt(20)
            run.font.bold = True
            run.font.color.rgb = RGBColor(27, 54, 93)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(14)
            i += 1
            continue
            
        if line.startswith('## '):
            p = doc.add_paragraph()
            run = p.add_run(line[3:].strip())
            run.font.name = 'Microsoft JhengHei'
            run.font.size = Pt(15)
            run.font.bold = True
            run.font.color.rgb = RGBColor(27, 54, 93)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(8)
            i += 1
            continue
            
        if line.startswith('### '):
            p = doc.add_paragraph()
            run = p.add_run(line[4:].strip())
            run.font.name = 'Microsoft JhengHei'
            run.font.size = Pt(12.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(43, 91, 132)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue
            
        if line.startswith('#### '):
            p = doc.add_paragraph()
            run = p.add_run(line[5:].strip())
            run.font.name = 'Microsoft JhengHei'
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(51, 78, 104)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        if line.startswith('> [!'):
            alert_type = 'NOTE'
            if 'IMPORTANT' in line: alert_type = 'IMPORTANT'
            elif 'WARNING' in line: alert_type = 'WARNING'
            elif 'TIP' in line: alert_type = 'TIP'
            
            callout_text = []
            i += 1
            while i < len(lines) and lines[i].startswith('>'):
                callout_text.append(lines[i].lstrip('>').strip())
                i += 1
                
            text_block = " ".join(callout_text)
            
            table = doc.add_table(rows=1, cols=1)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell = table.cell(0, 0)
            
            fill_color = "EBF8FF"
            title_text = "💡 重要提示"
            
            if alert_type == 'IMPORTANT':
                fill_color = "FEFCBF"
                title_text = "⚠️ 重點關注"
            elif alert_type == 'WARNING':
                fill_color = "FFF5F5"
                title_text = "🚨 合規警示"
            elif alert_type == 'TIP':
                fill_color = "F0FFF4"
                title_text = "✨ 趨勢焦點"

            set_cell_background(cell, fill_color)
            set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
            
            cp = cell.paragraphs[0]
            cp.paragraph_format.space_after = Pt(4)
            trun = cp.add_run(f"{title_text}: ")
            trun.font.name = 'Microsoft JhengHei'
            trun.font.bold = True
            trun.font.size = Pt(10)
            
            brun = cp.add_run(text_block)
            brun.font.name = 'Microsoft JhengHei'
            brun.font.size = Pt(10)
            
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
            continue
            
        if line.startswith('|') and '|' in line[1:]:
            table_lines = [line]
            i += 1
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i])
                i += 1
                
            rows_data = []
            for tline in table_lines:
                if '---' in tline:
                    continue
                cols = [c.strip() for c in tline.split('|')[1:-1]]
                if cols:
                    rows_data.append(cols)
                    
            if rows_data:
                table = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                for r_idx, row in enumerate(rows_data):
                    for c_idx, val in enumerate(row):
                        cell = table.cell(r_idx, c_idx)
                        cell.text = val
                        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
                        
                        p = cell.paragraphs[0]
                        p.paragraph_format.space_after = Pt(2)
                        for run in p.runs:
                            run.font.name = 'Microsoft JhengHei'
                            run.font.size = Pt(9)
                            
                        if r_idx == 0:
                            set_cell_background(cell, '1B365D')
                            for run in p.runs:
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(255, 255, 255)
                        else:
                            if r_idx % 2 == 1:
                                set_cell_background(cell, 'F7FAFC')
                            else:
                                set_cell_background(cell, 'FFFFFF')
                                
                doc.add_paragraph().paragraph_format.space_after = Pt(6)
            continue

        if line.startswith('* ') or line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.15
            
            content = line[2:].strip()
            parts = content.split('**')
            for idx, part in enumerate(parts):
                if not part: continue
                run = p.add_run(part)
                run.font.name = 'Microsoft JhengHei'
                run.font.size = Pt(10)
                if idx % 2 == 1:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(27, 54, 93)
            i += 1
            continue

        if line.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(line[2:].strip())
            run.font.name = 'Microsoft JhengHei'
            run.font.size = Pt(9.5)
            run.font.italic = True
            run.font.color.rgb = RGBColor(100, 116, 139)
            i += 1
            continue

        if line:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(5)
            p.paragraph_format.line_spacing = 1.15
            
            parts = line.split('**')
            for idx, part in enumerate(parts):
                if not part: continue
                run = p.add_run(part)
                run.font.name = 'Microsoft JhengHei'
                run.font.size = Pt(10)
                if idx % 2 == 1:
                    run.font.bold = True
            i += 1
        else:
            i += 1

    final_path = output_docx_path
    try:
        doc.save(final_path)
    except PermissionError:
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        base, ext = os.path.splitext(output_docx_path)
        final_path = f"{base}_{timestamp}{ext}"
        doc.save(final_path)
        print(f"[NOTICE] 原檔案已被鎖定，已另存為新檔：{final_path}")
        
    print(f"[SUCCESS] 精美 Word 報告已成功生成：{final_path}")
    return final_path

def parse_recipients(recipients_input):
    if isinstance(recipients_input, list):
        return [r.strip() for r in recipients_input if r.strip()]
    
    clean_str = recipients_input.replace('及', ';').replace('與', ';').replace(',', ';')
    return [r.strip() for r in clean_str.split(';') if r.strip()]

def send_via_outlook(to_emails, subject, body_html, attachment_path):
    try:
        import win32com.client
        outlook = win32com.client.Dispatch("Outlook.Application")
        mail = outlook.CreateItem(0)
        
        recipients_str = "; ".join(to_emails)
        mail.To = recipients_str
        mail.Subject = subject
        mail.HTMLBody = body_html
        if os.path.exists(attachment_path):
            mail.Attachments.Add(os.path.abspath(attachment_path))
        mail.Send()
        print(f"[SUCCESS] 已成功透過 Outlook 發送至多位收件人：{recipients_str}")
        return True
    except Exception as e:
        print(f"[DEBUG] Outlook COM 寄件測試失敗: {e}")
        return False

def send_email_with_attachment(to_emails, subject, body, attachment_path, smtp_server=None, smtp_port=None, from_email=None):
    recipients_list = parse_recipients(to_emails)
    
    # 優先嘗試本機 Outlook
    if send_via_outlook(recipients_list, subject, body, attachment_path):
        return True

    # 讀取環境變數或預設值
    env_smtp_server = os.environ.get("SMTP_SERVER", "localhost")
    env_smtp_port = int(os.environ.get("SMTP_PORT", 587 if os.environ.get("SMTP_USER") else 25))
    env_smtp_user = os.environ.get("SMTP_USER")
    env_smtp_pass = os.environ.get("SMTP_PASSWORD") or os.environ.get("SMTP_PASS")
    env_from_email = os.environ.get("SMTP_FROM") or "regulatory-service@bellwether-corp.com"

    actual_server = smtp_server or env_smtp_server
    actual_port = smtp_port or env_smtp_port
    actual_from = from_email or env_from_email
        
    msg = MIMEMultipart()
    msg['From'] = actual_from
    msg['To'] = "; ".join(recipients_list)
    msg['Subject'] = subject

    msg.attach(MIMEText(body, 'html', 'utf-8'))

    if os.path.exists(attachment_path):
        with open(attachment_path, "rb") as f:
            attach = MIMEApplication(f.read(), _subtype="docx")
            attach.add_header('Content-Disposition', 'attachment', filename=os.path.basename(attachment_path))
            msg.attach(attach)

    try:
        if actual_port == 465:
            server = smtplib.SMTP_SSL(actual_server, actual_port, timeout=15)
        else:
            server = smtplib.SMTP(actual_server, actual_port, timeout=15)
            if env_smtp_user:
                server.starttls()

        if env_smtp_user and env_smtp_pass:
            server.login(env_smtp_user, env_smtp_pass)

        server.sendmail(actual_from, recipients_list, msg.as_string())
        server.quit()
        print(f"[SUCCESS] 已透過 SMTP 成功寄出郵件給：{', '.join(recipients_list)}")
        return True
    except Exception as e:
        print(f"[STATUS_REPORT] 本機/雲端 SMTP 發送提示: {e}")
        print(f"[INFO] 報告 Word 檔已成功產出並存檔於：{attachment_path}")
        print(f"[INFO] 指定發送目標信箱：{', '.join(recipients_list)}")
        return False


def main():
    parser = argparse.ArgumentParser(description="跨國與台商內控內稽法規動態 Word 匯出與郵件寄送工具")
    parser.add_argument("--input", required=True, help="輸入的 Markdown 檔案路徑")
    parser.add_argument("--output", required=True, help="輸出的 Docx 檔案路徑")
    parser.add_argument("--email", nargs="?", const="", default="max.fanchiang@bellwether-corp.com; amelia.bui@bellwether-corp.com", help="收件人電子信箱（可填多個）")
    parser.add_argument("--title", default="最新台灣、越南(中越雙語)、泰國與中國大陸(含崑山台商)法規監理、內控內稽與合規因應指南週報", help="郵件主題與報告標題")
    
    args = parser.parse_args()
    raw_email = args.email.strip() if args.email else ""
    target_email = raw_email if raw_email else "max.fanchiang@bellwether-corp.com; amelia.bui@bellwether-corp.com"
    
    if not os.path.exists(args.input):
        print(f"[ERROR] 找不到輸入檔案：{args.input}")
        sys.exit(1)
        
    with open(args.input, 'r', encoding='utf-8') as f:
        md_content = f.read()
        
    actual_path = create_styled_docx(md_content, args.output)
    
    html_body = f"""
    <html>
    <body style="font-family: 'Microsoft JhengHei', sans-serif; color: #333;">
        <h2 style="color: #1B365D;">📋 {args.title}</h2>
        <p>您好：</p>
        <p>附件為最近 7 天針對<strong>台灣、越南 (Vietnam, 含中越雙語對照)、泰國 (Thailand) 及中國大陸 (含江蘇省/蘇州市/崑山地區台商規範)</strong> 最新新增與修訂法規之整理報告。報告特別全面強化新增<strong>【內部稽核 Internal Audit】與【內部控制制度 Internal Control System】</strong>專章與細則拆解。</p>
        <div style="background-color: #EBF8FF; border-left: 4px solid #3182CE; padding: 12px; margin: 16px 0;">
            <strong>📌 內部控制與內部稽核重點摘要：</strong>
            <ul>
                <li><strong>🇹🇼 台灣內控內稽辦法重大修訂 (2026/05 發布實施)</strong>：<br/>
                (1) 推動內部控制<strong>「三道模型」</strong>，要求自行查核由第二道督導以維持第三道（內部稽核）獨立性；<br/>
                (2) 強制設置<strong>法遵長、風管長與資安長「三長」</strong>；公發公司持續將「永續資訊」及「資安防護」列為年度必要內稽項目。</li>
                <li><strong>🇻🇳 越南內部稽核與內控實務 (Nghị định 05/2019 & 2026 審查)</strong>：<br/>
                <em>Thiết lập bộ phận kiểm toán nội bộ độc lập; kiểm tra tính tuân thủ đối với giá giao dịch liên kết FDI, hóa đơn điện tử và bảo vệ dữ liệu cá nhân (PDPD).</em></li>
                <li><strong>🇨🇳 中國大陸金稅四期與內控防線</strong>：大數據嚴查「四流一致」內控防線；崑山安全生產實行一票否決制。</li>
                <li><strong>🇹🇭 泰國 SEC COSO 內控自評</strong>：依 COSO 框架對內控五大要素自評並向審計委員會提出內部稽核報告。</li>
            </ul>
        </div>
        <p>詳細條款對照與企業落地因應指南請參閱附件中的精美 Word 報告排版內容。</p>
        <hr style="border: 0; border-top: 1px solid #eee;" />
        <p style="font-size: 12px; color: #777;">Antigravity 跨國與台商法規合規、內控內稽簡報系統自動發送</p>
    </body>
    </html>
    """
    
    send_email_with_attachment(target_email, args.title, html_body, actual_path)

if __name__ == "__main__":
    main()
